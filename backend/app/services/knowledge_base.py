"""
app/services/knowledge_base.py
Base de conocimiento de la empresa: extracción de texto de documentos
(PDF/Word/Markdown/texto plano), troceo en chunks y generación de embeddings.

Usado por routers/knowledge.py para procesar documentos subidos por el admin
y permitir que el chat IA responda preguntas institucionales (horarios,
políticas, sucursales, FAQs) — en paralelo al RAG existente sobre productos.
"""
import io
import logging
import re

from pypdf import PdfReader
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

ALLOWED_DOC_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/markdown": "md",
    "text/plain": "txt",
    "application/octet-stream": None,  # se determina por extensión
}
MAX_DOC_SIZE = 15 * 1024 * 1024  # 15 MB
MAX_PDF_PAGES = 200
MAX_EXTRACTED_CHARS = 250_000
MAX_CHUNKS = 250

# Tamaño aproximado de cada chunk (en caracteres) y solapamiento entre chunks
CHUNK_SIZE = 1200
CHUNK_OVERLAP = 150


def detect_file_type(filename: str, content_type: str | None) -> str | None:
    """Determina el tipo de archivo por content_type o, si no es claro, por extensión."""
    ct = (content_type or "").split(";")[0].strip().lower()
    if ct in ALLOWED_DOC_TYPES and ALLOWED_DOC_TYPES[ct]:
        return ALLOWED_DOC_TYPES[ct]

    ext = (filename or "").lower().rsplit(".", 1)[-1] if "." in (filename or "") else ""
    if ext in {"pdf", "docx", "md", "txt"}:
        return ext
    return None


def extract_text(file_bytes: bytes, file_type: str) -> str:
    """
    Extrae texto plano de un archivo según su tipo.
    Lanza ValueError si no se puede procesar.
    """
    try:
        if file_type == "pdf":
            reader = PdfReader(io.BytesIO(file_bytes))
            if len(reader.pages) > MAX_PDF_PAGES:
                raise ValueError(f"El PDF supera el máximo de {MAX_PDF_PAGES} páginas.")
            pages = []
            extracted_length = 0
            for page in reader.pages:
                page_text = page.extract_text() or ""
                extracted_length += len(page_text)
                if extracted_length > MAX_EXTRACTED_CHARS:
                    raise ValueError("El documento contiene demasiado texto para procesarlo.")
                pages.append(page_text)
            return "\n\n".join(pages).strip()

        if file_type == "docx":
            doc = DocxDocument(io.BytesIO(file_bytes))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            # Incluir también texto de tablas (horarios suelen venir en tablas)
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            text = "\n".join(parts).strip()
            if len(text) > MAX_EXTRACTED_CHARS:
                raise ValueError("El documento contiene demasiado texto para procesarlo.")
            return text

        if file_type in ("md", "txt"):
            text = file_bytes.decode("utf-8").strip()
            if len(text) > MAX_EXTRACTED_CHARS:
                raise ValueError("El documento contiene demasiado texto para procesarlo.")
            return text

    except Exception as e:
        logger.error(f"Error extrayendo texto ({file_type}): {e}")
        raise ValueError(f"No se pudo leer el archivo {file_type.upper()}. ¿Está dañado o protegido?")

    raise ValueError(f"Tipo de archivo no soportado: {file_type}")


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """
    Trocea el texto en fragmentos manejables para embeddings, intentando
    cortar en límites de párrafo/oración para no partir ideas a la mitad.
    """
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not text:
        return []

    if len(text) <= chunk_size:
        return [text]

    chunks: list[str] = []
    start = 0
    n = len(text)

    while start < n:
        end = min(start + chunk_size, n)

        if end < n:
            # Buscar el último salto de párrafo o punto dentro de la ventana
            window = text[start:end]
            cut = max(window.rfind("\n\n"), window.rfind(". "), window.rfind("\n"))
            if cut > chunk_size * 0.5:  # solo si el corte no es demasiado temprano
                end = start + cut + 1

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
            if len(chunks) > MAX_CHUNKS:
                raise ValueError("El documento genera demasiados fragmentos para procesarlo.")

        if end >= n:
            break
        start = max(end - overlap, start + 1)

    return chunks
